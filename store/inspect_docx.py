import sys
from pathlib import Path

from docx import Document


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    if len(sys.argv) < 2:
        raise SystemExit("usage: inspect_docx.py <docx>")
    path = Path(sys.argv[1])
    doc = Document(path)
    print(f"FILE: {path}")
    print(f"PARAS {len(doc.paragraphs)} TABLES {len(doc.tables)}")
    print("---PARAGRAPHS---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"{i:03d}: {text}")
    print("---TABLES---")
    for ti, table in enumerate(doc.tables):
        print(f"TABLE {ti}")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            print(f"{ri:03d}: " + " | ".join(cells))


if __name__ == "__main__":
    main()
