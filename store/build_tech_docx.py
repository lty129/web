from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "商城项目技术总结.md"
OUTPUT = ROOT / "商城项目技术总结.docx"

ACCENT = "23464A"
MUTED = "6F7977"
LIGHT = "F4F1EA"
BORDER = "D9DED9"


def set_east_asia_font(run, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("商城项目技术总结")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia_font(r)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    fp = section.footer.paragraphs[0]
    fp.text = ""
    r = fp.add_run("第 ")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia_font(r)
    add_page_number(fp)
    r = fp.add_run(" 页")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia_font(r)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1F2625")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Title", 26, ACCENT, 0, 8),
        ("Subtitle", 11, MUTED, 0, 14),
        ("Heading 1", 16, ACCENT, 14, 6),
        ("Heading 2", 13, "1F2625", 10, 4),
        ("Heading 3", 11, "1F2625", 8, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if "Heading" in style_name or style_name == "Title":
            style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_metadata_table(doc):
    table = doc.add_table(rows=4, cols=2)
    set_table_width(table)
    table.style = "Table Grid"
    rows = [
        ("项目名称", "Maison Store 高级生活方式商城"),
        ("文档用途", "期末项目答辩技术总结"),
        ("核心技术", "HTML、CSS、JavaScript、Node.js、MySQL 8.0"),
        ("运行地址", "http://localhost:5173"),
    ]
    for row, (key, value) in zip(table.rows, rows):
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(4.9)
        for idx, text in enumerate((key, value)):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            if idx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.bold = idx == 0
            run.font.color.rgb = RGBColor.from_string(ACCENT if idx == 0 else "1F2625")
            set_east_asia_font(run)
    doc.add_paragraph()


def add_code_block(doc, lines):
    text = "\n".join(lines).strip("\n")
    if not text:
        return
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F6F2")
    set_cell_border(cell, "E1E4DF", "4")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9.2)
        run.font.color.rgb = RGBColor.from_string("35413F")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    set_east_asia_font(run)


def add_heading(doc, text, level):
    style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_east_asia_font(run)


def build_docx():
    md = SOURCE.read_text(encoding="utf-8").replace("\r\n", "\n")
    lines = md.splitlines()
    title = lines[0].lstrip("#").strip() if lines and lines[0].startswith("#") else "商城项目技术总结"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    configure_styles(doc)
    add_header_footer(doc)

    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title_p.add_run(title)
    set_east_asia_font(tr)

    sub = doc.add_paragraph(style="Subtitle")
    sr = sub.add_run("Maison Store 商城系统 | 前端、后端、数据库与部署说明")
    set_east_asia_font(sr)
    add_metadata_table(doc)

    in_code = False
    code_lines = []
    skip_first_title = True
    for raw in lines[1:]:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("# "):
            if skip_first_title:
                skip_first_title = False
            else:
                add_heading(doc, line[2:].strip(), 2)
        else:
            add_body_paragraph(doc, line.strip())

    if code_lines:
        add_code_block(doc, code_lines)

    doc.core_properties.title = title
    doc.core_properties.subject = "期末答辩技术总结"
    doc.core_properties.author = "Maison Store 项目组"
    doc.core_properties.keywords = "商城, HTML, CSS, JavaScript, Node.js, MySQL"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
